import re

import markdown
from bs4 import BeautifulSoup
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Twips
from docxtpl import RichText as BaseRichText

FONT = 'Noto Sans'
# Bullet glyph per nesting depth (level 1 outermost): filled circle, then
# hollow circle, then square for level 3 and deeper.
BULLET_LEVELS = ['●', '○', '■']
# Invisible markers prefixed to heading paragraphs so the post-render step
# can identify them and apply paragraph-level spacing.
HEADING_MARKERS = {
    'h1': '\u200b\u200bH1\u200b\u200b',
    'h2': '\u200b\u200bH2\u200b\u200b',
    'h3': '\u200b\u200bH3\u200b\u200b',
    'h4': '\u200b\u200bH4\u200b\u200b',
    'h5': '\u200b\u200bH5\u200b\u200b',
    'h6': '\u200b\u200bH6\u200b\u200b',
}


class RichText(BaseRichText):
    """
    Fixed RichText with correct OOXML element ordering.
    Overrides add() to put rFonts BEFORE b/i formatting elements.
    """

    def add(
        self,
        text,
        style=None,
        color=None,
        highlight=None,
        size=None,
        subscript=None,
        superscript=None,
        bold=False,
        italic=False,
        underline=False,
        strike=False,
        font=None,
        url_id=None,
        rtl=False,
        lang=None,
    ):
        from xml.sax.saxutils import escape

        # If a RichText is added
        if isinstance(text, RichText):
            self.xml += text.xml
            return

        # If not a string: cast to string
        if not isinstance(text, (str, bytes)):
            text = str(text)
        if not isinstance(text, str):
            text = text.decode("utf-8", errors="ignore")
        text = escape(text)

        prop = ""

        # OOXML spec order: rStyle, rFonts, b, bCs, i, iCs, ...
        if style:
            prop += '<w:rStyle w:val="%s"/>' % style
        if font:  # MOVED: rFonts must come before b/i
            regional_font = ""
            if ":" in font:
                region, font = font.split(":", 1)
                regional_font = ' w:{region}="{font}"'.format(font=font, region=region)
            prop += '<w:rFonts w:ascii="{font}" w:hAnsi="{font}" w:cs="{font}"{regional_font}/>'.format(
                font=font, regional_font=regional_font
            )
        if bold:
            prop += "<w:b/>"
            if rtl:
                prop += "<w:bCs/>"
        if italic:
            prop += "<w:i/>"
            if rtl:
                prop += "<w:iCs/>"
        if strike:
            prop += "<w:strike/>"
        if color:
            if color[0] == "#":
                color = color[1:]
            prop += '<w:color w:val="%s"/>' % color
        if size:
            prop += '<w:sz w:val="%s"/>' % size
            prop += '<w:szCs w:val="%s"/>' % size
        if highlight:
            if highlight[0] == "#":
                highlight = highlight[1:]
            prop += '<w:shd w:fill="%s"/>' % highlight
        if underline:
            if underline not in [
                "single", "double", "thick", "dotted", "dash",
                "dotDash", "dotDotDash", "wave",
            ]:
                underline = "single"
            prop += '<w:u w:val="%s"/>' % underline
        if subscript:
            prop += '<w:vertAlign w:val="subscript"/>'
        if superscript:
            prop += '<w:vertAlign w:val="superscript"/>'
        if rtl:
            prop += '<w:rtl w:val="true"/>'
        if lang:
            prop += '<w:lang w:val="%s"/>' % lang

        xml = "<w:r>"
        if prop:
            xml += "<w:rPr>%s</w:rPr>" % prop
        xml += '<w:t xml:space="preserve">%s</w:t></w:r>' % text
        if url_id:
            xml = '<w:hyperlink r:id="%s" w:tgtFrame="_blank">%s</w:hyperlink>' % (
                url_id,
                xml,
            )
        self.xml += xml

### REPLACE THIS: RMV
### https://stackoverflow.com/questions/70363269/how-can-i-convert-a-markdown-string-to-a-docx-in-python
### https://stackoverflow.com/questions/51829366/bullet-lists-in-python-docx/51830413#51830413

def list_number(doc, par, style, prev=None, level=None, num=True):
    """
    Makes a paragraph into a list item with a specific level and
    optional restart.

    An attempt will be made to retreive an abstract numbering style that
    corresponds to the style of the paragraph. If that is not possible,
    the default numbering or bullet style will be used based on the
    ``num`` parameter.

    Parameters
    ----------
    doc : docx.document.Document
        The document to add the list into.
    par : docx.paragraph.Paragraph
        The paragraph to turn into a list item.
    prev : docx.paragraph.Paragraph or None
        The previous paragraph in the list. If specified, the numbering
        and styles will be taken as a continuation of this paragraph.
        If omitted, a new numbering scheme will be started.
    level : int or None
        The level of the paragraph within the outline. If ``prev`` is
        set, defaults to the same level as in ``prev``. Otherwise,
        defaults to zero.
    num : bool
        If ``prev`` is :py:obj:`None` and the style of the paragraph
        does not correspond to an existing numbering style, this will
        determine wether or not the list will be numbered or bulleted.
        The result is not guaranteed, but is fairly safe for most Word
        templates.
    """
    xpath_options = {
        True: {'single': 'count(w:lvl)=1 and ', 'level': 0},
        False: {'single': '', 'level': level},
    }

    def style_xpath(prefer_single=True):
        """
        The style comes from the outer-scope variable ``par.style.name``.
        """
        return (
            'w:abstractNum['
                '{single}w:lvl[@w:ilvl="{level}"]/w:pStyle[@w:val="{style}"]'
            ']/@w:abstractNumId'
        ).format(style=style, **xpath_options[prefer_single])

    def type_xpath(prefer_single=True):
        """
        The type is from the outer-scope variable ``num``.
        """
        type = 'decimal' if num else 'bullet'
        return (
            'w:abstractNum['
                '{single}w:lvl[@w:ilvl="{level}"]/w:numFmt[@w:val="{type}"]'
            ']/@w:abstractNumId'
        ).format(type=type, **xpath_options[prefer_single])
    def get_abstract_id():
        """
        Select as follows:

            1. Match single-level by style (get min ID)
            2. Match exact style and level (get min ID)
            3. Match single-level decimal/bullet types (get min ID)
            4. Match decimal/bullet in requested level (get min ID)
            3. 0
        """
        for fn in (style_xpath, type_xpath):
            for prefer_single in (True, False):
                xpath = fn(prefer_single)
                ids = numbering.xpath(xpath)
                if ids:
                    return min(int(x) for x in ids)
        return 0

    if (prev is None or
            prev._p.pPr is None or
            prev._p.pPr.numPr is None or
            prev._p.pPr.numPr.numId is None):
        if level is None:
            level = 0
        numbering = doc.part.numbering_part.numbering_definitions._numbering
        # Compute the abstract ID first by style, then by num
        anum = get_abstract_id()
        # Set the concrete numbering based on the abstract numbering ID
        numbr = numbering.add_num(anum)
        # Make sure to override the abstract continuation property
        numbr.add_lvlOverride(ilvl=level).add_startOverride(1)
        # Extract the newly-allocated concrete numbering ID
        numbr = numbr.numId
    else:
        if level is None:
            level = prev._p.pPr.numPr.ilvl.val
        # Get the previous concrete numbering ID
        numbr = prev._p.pPr.numPr.numId.val
    if num: 
        par._p.get_or_add_pPr().get_or_add_numPr().get_or_add_numId().val = numbr
    par._p.get_or_add_pPr().get_or_add_numPr().get_or_add_ilvl().val = level

def apply_list_indentation(doc):
    """Post-render: add hanging indent to bullet/numbered list paragraphs.

    Detects leading tabs (used for nesting level), strips them from the
    rendered text, and replaces them with proper paragraph indentation
    so wrapped lines align with the content after the bullet.

    Args:
        doc: A DocxTemplate instance (after render).
    """
    pattern = re.compile(r'^(\t+)([●○■]|\d+\.)\t')

    def _fix(paragraphs):
        for para in paragraphs:
            match = pattern.match(para.text)
            if match:
                level = len(match.group(1))

                # Strip the leading level-tabs from the actual runs;
                # paragraph indentation handles nesting instead.
                tabs_remaining = level
                for run in para.runs:
                    while tabs_remaining > 0 and run.text.startswith('\t'):
                        run.text = run.text[1:]
                        tabs_remaining -= 1
                    if tabs_remaining == 0:
                        break

                # After stripping, text is "●\t<content>".
                # left_indent  = where content + wrapped text sits
                # hanging (negative first_line_indent) = space for the bullet
                para.paragraph_format.left_indent = Twips(level * 720)
                para.paragraph_format.first_line_indent = Twips(-360)

    _fix(doc.docx.paragraphs)
    for table in doc.docx.tables:
        for row in table.rows:
            for cell in row.cells:
                _fix(cell.paragraphs)


def fix_invalid_tables(doc):
    """Post-render: repair tables left structurally invalid by missing data.

    Args:
        doc: A DocxTemplate instance (after render).
    """
    for tbl in list(doc.docx.element.iter(qn('w:tbl'))):
        parent = tbl.getparent()
        if parent is None:
            continue
        # Empty table (no rows) -> swap for an empty paragraph (always valid,
        # never leaves two tables adjacent or a trailing table).
        if tbl.find(qn('w:tr')) is None:
            parent.replace(tbl, OxmlElement('w:p'))
            continue
        # Every cell must contain at least one paragraph (or nested table).
        for tc in tbl.iter(qn('w:tc')):
            if tc.find(qn('w:p')) is None and tc.find(qn('w:tbl')) is None:
                tc.append(OxmlElement('w:p'))


def apply_heading_spacing(doc):
    """Post-render: apply paragraph spacing to heading paragraphs.

    Detects zero-width-space markers left by mark2html, strips them,
    and applies space-before to the paragraph.

    Args:
        doc: A DocxTemplate instance (after render).
    """
    # Build a lookup: marker text -> heading level config
    # space_before in points for each heading level
    heading_config = {
        'h1': {'space_before': 12, 'space_after': 4},
        'h2': {'space_before': 10, 'space_after': 4},
        'h3': {'space_before': 8, 'space_after': 2},
        'h4': {'space_before': 6, 'space_after': 2},
        'h5': {'space_before': 4, 'space_after': 2},
        'h6': {'space_before': 4, 'space_after': 2},
    }
    marker_pattern = re.compile(r'\u200b\u200b(H[1-6])\u200b\u200b')

    def _fix(paragraphs):
        for para in paragraphs:
            if not para.runs:
                continue
            first_run = para.runs[0]
            match = marker_pattern.match(first_run.text)
            if match:
                tag = match.group(1).lower()
                first_run.text = first_run.text[match.end():]
                config = heading_config.get(tag, {})
                if 'space_before' in config:
                    para.paragraph_format.space_before = Pt(config['space_before'])
                if 'space_after' in config:
                    para.paragraph_format.space_after = Pt(config['space_after'])

    _fix(doc.docx.paragraphs)
    for table in doc.docx.tables:
        for row in table.rows:
            for cell in row.cells:
                _fix(cell.paragraphs)


def mark2html(value, font_size=None):
    if font_size is not None:
        # OOXML font size is in half-points, so multiply by 2
        font_size = int(font_size) * 2
    if value is None or not isinstance(value, str):
        return []
    # Fix corrupted HTML entity: &#191; (¿) appears where apostrophes should be
    value = value.replace('&#191;', "'")
    html = markdown.markdown(value)
    soup = BeautifulSoup(html, features='html.parser')

    def nested(children, level=0, numbered=False, font_size=None):
        list_index = 0
        for n, tag in enumerate(children):
            if tag.name == 'li':
                list_index += 1
                indent = "".join("\t" for _ in range(level))
                if numbered:
                    indent += f"{list_index}.\t"
                else:
                    bullet = BULLET_LEVELS[min(max(level - 1, 0), len(BULLET_LEVELS) - 1)]
                    indent += f"{bullet}\t"
                rt = RichText()
                rt.add(indent, font=FONT, size=font_size)

                # Check if this <li> contains nested lists
                nested_lists = [c for c in tag.children if hasattr(c, 'name') and c.name in ('ul', 'ol')]

                if nested_lists:
                    # <li> has nested lists — extract inline text separately,
                    # then yield nested list items at a deeper level.
                    # parseHtmlToDoc handles text and inline elements and
                    # naturally skips <ul>/<ol> children.
                    para = parseHtmlToDoc(tag, level, list_index if numbered else False, font_size=font_size)
                    if para:
                        yield [rt] + para
                    for nl in nested_lists:
                        yield from nested(nl.children, level=level+1, numbered=(nl.name == 'ol'), font_size=font_size)
                else:
                    # No nested lists — use standard logic
                    following = iter(nested(tag.children, level, numbered, font_size))
                    first = next(following, None)
                    if first is not None:
                        row = [rt] + first
                        yield row
                        yield from following
                    else:
                        # li contains direct text with no nested block elements
                        para = parseHtmlToDoc(tag, level, list_index if numbered else False, font_size=font_size)
                        row = [rt] + para
                        yield row
            elif tag.name in ('p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6'):
                para = parseHtmlToDoc(tag, level, n if numbered else False, font_size=font_size)
                if tag.name in HEADING_MARKERS:
                    marker_rt = RichText()
                    marker_rt.add(HEADING_MARKERS[tag.name], font=FONT, size=2)
                    para = [marker_rt] + para
                yield para
            elif tag.name in ('ul', 'ol'):
                yield from nested(tag.children, level=level+1, numbered=(tag.name == 'ol'), font_size=font_size)
    paragraphs = [
        tag for tag in nested(soup, font_size=font_size)
        if tag and any(
            not isinstance(rt, RichText) or rt.xml.strip()
            for rt in tag
        )
    ]

    return paragraphs

def parseHtmlToDoc(org_tag, level=0, numbered=False, font_size=None):
    contents = org_tag.contents
    pars = []
    for con in contents:
        if str(type(con)) == "<class 'bs4.element.Tag'>":
            tag = con
            if tag.name == 'br':
                # Add a line break - create new RichText with break XML
                if len(pars) > 0 and isinstance(pars[-1], RichText):
                    pars[-1].xml += '<w:r><w:br/></w:r>'
                else:
                    source = RichText()
                    source.xml = '<w:r><w:br/></w:r>'
                    pars.append(source)
            elif tag.name in ('strong',"h1","h2","h3","h4","h5","h6"):
                # Append bold text to existing RichText or create new one
                if len(pars) > 0 and isinstance(pars[-1], RichText):

                    pars[-1].add(con.contents[0], bold=True, font=FONT, size=font_size)
                else:
                    source = RichText()
                    source.add(con.contents[0], bold=True, font=FONT, size=font_size)
                    pars.append(source)
            elif tag.name == 'img':
                # Inline images in markdown blobs aren't supported here: this
                # parser has no DocxTemplate reference and InlineImage requires
                # one. Templates that need an image should use a dedicated
                # context key + the to_image filter instead.
                continue
            elif tag.name == 'em':
                # Append italic text to existing RichText or create new one
                if len(pars) > 0 and isinstance(pars[-1], RichText):
                    pars[-1].add(con.contents[0], italic=True, font=FONT, size=font_size)
                else:
                    source = RichText()
                    source.add(con.contents[0], italic=True, font=FONT)
                    pars.append(source)
            elif tag.name in ('ul', 'ol'):
                # Nested lists are emitted separately by nested(); skip them
                # here so their text isn't duplicated into the parent li.
                continue
            else:
                # Unrecognised tag (e.g. <http:> from unescaped URLs) —
                # recurse into its children so nested content isn't lost.
                pars.extend(parseHtmlToDoc(tag, level, numbered, font_size=font_size))
        else:
            # Regular text - collapse insignificant HTML whitespace (source
            # newlines/indentation) to single spaces like a browser would, and
            # drop it at the paragraph start. Otherwise a literal newline inside
            # an <li> renders as a blank line between bullets. Only ASCII
            # whitespace is collapsed/stripped so authored non-breaking spaces
            # (\xa0) survive into the docx.
            con = re.sub(r'[ \t\n\r\f\v]+', ' ', str(con))
            if not pars:
                con = con.lstrip(' \t\n\r\f\v')
            if not con:
                continue
            # append to existing RichText or create new one
            if len(pars) > 0 and isinstance(pars[-1], RichText):
                pars[-1].add(con, font=FONT, size=font_size)
            else:
                source = RichText()
                if org_tag.name == 'h2':
                    source.add(con, bold=True, size=font_size or 40, font=FONT)
                else:
                    source.add(con, font=FONT, size=font_size)
                pars.append(source)
    return pars
